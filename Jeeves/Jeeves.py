import subprocess                   #Work with computer processes and executables
import pathlib                      #Work with directory paths on local
import speech_recognition as sr     #Speech recognition lib
from docx import Document           #Word Doc lib
import datetime as dt               #Gather today's date
import pyowm                        #Weather API lib
import time                         #Grab current time
import requests                     #HTTP requests lib
import json                         #Parsing JSONs from Spotify API
import base64                       #Convert to Base64 String for Spotify API passcode
import ast                          #Turns strings to dictionaries
import wikipedia                    #Wiki scraper package
import pyttsx3                      #Text to speech response package
import playsound                    #Capable of playing sounds
from newsapi import NewsApiClient   #News headlines API lib

owAPI = #Insert your API key to Open Weather Maps here
spotifyID = #Insert Spotify Client ID here
spotifySecret = #Insert Spotify Client Secret here
spotifyRefreshToken =
spotifyRefreshHeader =
newsAPIkey = #Insert your API key to NewsAPI here

#Initialise the News API
newsapi = NewsApiClient(newsAPIkey)

class TTS:
    engine = None
    rate = None

    #Set initialisation properties for the voice here
    def __init__(self):
        self.engine = pyttsx3.init() #Initialise text to speech functionality
        self.engine.setProperty('voice',self.engine.getProperty('voices')[7].id) #Set British Male voice
        self.engine.setProperty('rate',250)

    #Instance method for speaking aloud
    def speak(self, text):
        self.engine.say(text)
        self.engine.runAndWait()

class Commands:
    tts = TTS() #Define prefix for calling tts

    @classmethod
    def openApp(cls, app, rootname="none"):
        """
        Function which paths to Apps directory and opens an app using subprocess based on user input
        Inputs: app - the name (exactly as given in command line) of the app to open (str)
                rootname (optional) - if the App's name is different to its containing directory on the mac, rootname is the actual executable's name
                e.g. Microsoft Teams.app/Contents/MacOS/Teams where Microsoft Teams.app is the app and Teams is the rootname executable
        """

        if rootname == "none":
            subprocess.Popen(["../../../../Applications/{0}.app/Contents/MacOS/{0}".format(app),"-i"])
        else:
            subprocess.Popen(["../../../../Applications/{0}.app/Contents/MacOS/{1}".format(app,rootname),"-i"])

    @classmethod
    def dictate(cls, notes):
        """
        Function which takes dictated input and creates a word doc filled with the dictated text
        Input: notes - the recorded input to be transcribed (str)
        """

        doc = Document()     #Create doc using constructor
        today = dt.date.today() #Take date for heading

        #List to split paragraphs
        paras = []

        #Replacers to change ends of sentences to ". "
        while "end sentence" in notes or "and sentence" in notes:
            loc = notes.find("end sentence")
            if loc != -1:
                notes[loc+14].capitalize()
            loc2 = notes.find("and sentence")
            if loc2 != -1:
                notes[loc2+14].capitalize()
            notes = notes.replace(" end sentence", ".")
            notes = notes.replace(" and sentence", ".")

        #Replacers allow user to escape the words 'paragraph, paragraphs' if included in notes
        while "escape paragraph" in notes:
            notes = notes.replace("escape paragraph", "pescape")
        while "paragraphs" in notes:
            notes = notes.replace("paragraphs", "psescape")

        while "paragraph" in notes:
            loc = notes.index("paragraph")
            current = notes[:loc] #Split into next paragraph at keyword "paragraph"

            while "pescape" in current:
                current = current.replace("pescape","paragraph")
            while "psescape" in current:
                current = current.replace("psescape", "paragraphs")

            paras.append(current)
            notes = notes[loc+10:] #Cut appended paragraph from notes

        while "pescape" in notes:
            notes = notes.replace("pescape","paragraph")
        while "psescape" in notes:
            notes = notes.replace("psescape", "paragraphs")
        paras.append(notes)

        #Check if a similarly named doc already exists, save to new
        notestoday = []
        for i in pathlib.Path("Dictations").iterdir():
            if str(today) in str(i):
                notestoday.append(str(i))
        name = "Dictations/{0}-{1}.docx".format(today,len(notestoday)+1)

        doc.add_heading("Notes for {0} (Entry {1})".format(today,len(notestoday)+1), 0)

        for i in paras:
            doc.add_paragraph(i)  #Add notes to document
        
        doc.save(name)
        print("All done sir! You'll find it under the name {0}.".format(name))
        Commands.tts.speak("All done sir! You'll find it under the name {0}.".format(name))

    @classmethod
    def weather(cls):
        """
        Function that returns the weather data for the current time grabbed from the OpenWeatherMap API
        """

        owm = pyowm.OWM(owAPI)
        observation = owm.weather_at_place("London,UK") #Grab weather object from API

        w = observation.get_weather() #Pull weather from object
        temp = w.get_temperature("celsius") #Temperature dict in celsius
        wind = w.get_wind("miles_hour") #Wind dict
        rain = w.get_rain() #Rain dict

        print("The temperature is currently {0}˚C, ranging between {1}˚C and {2}˚C.".format(temp["temp"],temp["temp_max"],temp["temp_min"]))
        Commands.tts.speak("The temperature is currently {0}˚C, ranging between {1}˚C and {2}˚C.".format(temp["temp"],temp["temp_max"],temp["temp_min"]))

        if len(rain) == 0:
            print("There's no rain at the moment.")
            Commands.tts.speak("There's no rain at the moment.")


        print("The wind is currently blowing at {0:0.2f}mph on bearing {1}˚.".format(wind["speed"],wind["deg"]))
        Commands.tts.speak("The wind is currently blowing at {0:0.2f}mph on bearing {1}˚.".format(wind["speed"],wind["deg"]))

    @classmethod
    def forecast(cls, hr):
        """
        Function that returns a weather prediction by grabbing data from the OpenWeatherMap API
        Input: hr - the hour to forecast for, must be 1, 2 or 3 (three hours max) (int)
        """

        if hr > 3:
            print("Sorry sir, I can only forecast over the next three hours.")
            Commands.tts.speak("Sorry sir, I can only forecast over the next three hours.")

            return None

        owm = pyowm.OWM(owAPI)
        forecast = owm.three_hours_forecast("London,UK") #Grab weather forecast object from API

        #Format time to ISO8601-formatted string to grab correct data
        today = dt.date.today()
        hour = time.gmtime(time.time()+3600+(hr*3600)).tm_hour
        minute = time.gmtime(time.time()+3600+(hr*3600)).tm_min
        second = time.gmtime(time.time()+3600+(hr*3600)).tm_sec
        ftime = "{0} {1}:{2}:{3}+00".format(today,hour,minute,second)

        fw = forecast.get_weather_at(ftime) #Pull weather from forecast object
        ftemp = fw.get_temperature("celsius") #Temperature dict in celsius
        fwind = fw.get_wind("miles_hour") #Wind dict in mph
        frain = fw.get_rain() #Rain dict

        print("The temperature in {0} hours' time will range between {1}˚C and {2}˚C.".format(hr,ftemp["temp_min"],ftemp["temp_max"]))
        Commands.tts.speak("The temperature in {0} hours' time will range between {1}˚C and {2}˚C.".format(hr,ftemp["temp_min"],ftemp["temp_max"]))
    
        if hr == 1: 
            if "1h" not in frain: #No rain predicted in next hour
                print("No rain is forecast for the next hour.")
                Commands.tts.speak("No rain is forecast for the next hour.")
    
            else:
                print("{0}mm of rain is expected over the next hour.".format(frain["1h"]))
                Commands.tts.speak("{0}mm of rain is expected over the next hour.".format(frain["1h"]))
    

        elif hr == 2:
            if "2h" not in frain:
                print("No rain is forecast for the next two hours.")
                Commands.tts.speak("No rain is forecast for the next two hours.")
    
            else:
                print("{0}mm of rain is expected over the next two hours.".format(frain["2h"]))
                Commands.tts.speak("{0}mm of rain is expected over the next two hours.".format(frain["2h"]))
    

        elif hr == 3:
            if "3h" not in frain:
                print("No rain is forecast for the next three hours.")
                Commands.tts.speak("No rain is forecast for the next three hours.")
    
            else:
                print("{0}mm of rain is expected over the next three hours.".format(frain["3h"]))
                Commands.tts.speak("{0}mm of rain is expected over the next three hours.".format(frain["3h"]))
    


        print("There will be a {0:0.2f}mph wind on bearing {1}˚.".format(fwind["speed"],fwind["deg"]))
        Commands.tts.speak("There will be a {0:0.2f}mph wind on bearing {1}˚.".format(fwind["speed"],fwind["deg"]))

    files = [] #Necessary class variable for completelist/filelist

    @classmethod
    def completelist(cls, d):

        for i in pathlib.Path(d).iterdir():
            if (".DS_Store" in str(i) or "__pycache__" in str(i) or ".ipynb_checkpoints" in str(i)):
                continue
            else:
                if str(i) not in Commands.files:
                    Commands.files.append(str(i))
                else:
                    pass
            try:
                Commands.completelist(i)
            except NotADirectoryError:
                continue
            
        f = [] #Local variable to return so class variable can be cleared
        for i in Commands.files:
            f.append(i)

        return f

    @classmethod
    def filelist(cls, d, search):
    
        files = Commands.completelist(d)
        Commands.files = [] #Clear files class variable for re-use
        results = []
    
        for i in files:
            if search in i:
                results.append(i)
            else:
                continue
        
        if len(results) == 0:
            print("I'm afraid I couldn't find any files with that name in sir. Sorry!")
            Commands.tts.speak("I'm afraid I couldn't find any files with that name in sir. Sorry!")

        else:
            print("I've found the following list of files that might fit your purpose sir:\n{0}".format(results))
            Commands.tts.speak("I've found the following list of files that might fit your purpose sir")

    access = "" #Necessary class variable taking place of Spotify API Token

    @classmethod
    def refreshSpotifyToken(cls, sRT, sRH):
        """
        Function that pipes to command line to curl request a new spotify token
        Inputs: sRT - spotifyRefreshToken (str)
                sRH - spotifyRefreshHeader (str)
        Defines 'access' class variable, currently active auth token for Spotify API
        """
        #Pipe subprocess command to terminal, curl request refresh token
        get_token = subprocess.Popen(
            ["curl", "-H", sRH, "-d", "grant_type=refresh_token", "-d", "refresh_token="+sRT, "https://accounts.spotify.com/api/token"],
            stdout = subprocess.PIPE,
        )
        out = get_token.communicate() #Returns json object of Spotify API data

        #Stringify byte iformation to slice
        out = str(out)
        opening = out.find("{") #Start of object
        closing = out.find("}") #Close of object
        out = out[opening:closing+1]

        out = ast.literal_eval(out) #Turn string into dictionary
        Commands.access = out["access_token"] #Access token for Spotify data

    @classmethod
    def define(cls, topic, sentences=4):
        try:
            wiki = wikipedia.summary(topic, sentences=sentences)
            print("The topic you searched for came out as '{0}'".format(topic))
            print("Here's what I found sir:")
            Commands.tts.speak("The topic you searched for came out as '{0}'. Here's what I found sir.".format(topic))
            print(wiki)
            Commands.tts.speak(wiki)
        except:
            print("Sorry sir, I ran into some errors trying to fetch that. Feel free to try again.")
            Commands.tts.speak("Sorry sir, I ran into some errors trying to fetch that. Feel free to try again.")

    @classmethod
    def episode(cls):
        #Fetch correct episode number from separate file
        #As file updates in real time but Jeeves cannot, must be saved separately
        with open("episodecounter.py") as f:
            line = str(f.readlines()[0])
            starTrekEpisode = int(line[18:]) #Trim for number

        #Open csv with episodes
        with open("startrekepisodes.csv", encoding="utf-8") as f:
            line = f.readlines()[starTrekEpisode] #Read line attributed to correct episode
            data = []
            while len(str(line)) > 0: #Separate line string into list of entries
                loc = line.find(",")
                if loc != -1:
                    entry = line[:loc]
                    data.append(entry)
                    line = line[loc+1:]
                else:
                    data.append(line)
                    line = ""

        #Current episode data
        season = data[2]
        episode = data[3]
        episode = "Season " + episode #Add to string for text to speech
        episode = episode.replace("x"," Episode ")

        #Say current episode
        print("The next episode to watch is {0}, {1}.".format(season,episode))
        Commands.tts.speak("The next episode to watch is {0}, {1}.".format(season,episode))

        #Add one to episode (moving to next episode) and overwrite file
        content = ["starTrekEpisode = {n}".format(n=starTrekEpisode+1)]
        with open("episodecounter.py", "w") as f:
            f.write(content[0])

    @classmethod
    def cast(cls, topic):
        text = wikipedia.WikipediaPage(topic).html() #Grab html from wikipedia

        #Trim html tags for list of starring
        text = text[text.index("Starring"):]
        openlist = text.index("<ul>")
        closelist = text.index("</ul>")
        text = text[openlist:closelist]

        #New list to fill with cast names
        cast = []
        while len(text) > 0:
            openitem = text.find('">')  #Trim list tags to get names
            closeitem = text.find("</a>")
            if openitem != -1:
                name = text[openitem+2:closeitem]
                cast.append(name)
            else:
                text = ""
            text = text[closeitem+4:]

        #Output to terminal and speak aloud
        print("The cast of {0} includes".format(topic))
        Commands.tts.speak("The cast of {0} includes".format(topic))
        print(cast)
        for i in cast:
            Commands.tts.speak(i)

    @classmethod
    def news(cls):
        top_headlines = newsapi.get_top_headlines(language='en',sources='bbc-news')
        print("Here are the headlines from BBC News:")
        Commands.tts.speak("Here are the headlines from BBC News.")
        for i in top_headlines["articles"]:
            print("-> {0}".format(i["title"]))
            Commands.tts.speak(i["title"])

def Jeeves():
    tts = TTS() #Define prefix for calling tts

    r = sr.Recognizer()
    mic = sr.Microphone()

    attending = True

    """if time.gmtime().tm_hour > 11: #Check it's afternoon
        print("Good afternoon sir, what can I do for you today?")
        tts.speak("Good afternoon sir, what can I do for you today?")
    else:
        print("Good morning sir, what can I do for you today?")
        tts.speak("Good morning sir, what can I do for you today?")"""

    while attending == True:

        playsound.playsound("jeeves_sound.mp3")

        with mic as source:
            audio = r.listen(source, phrase_time_limit=7.0)

        prompt = r.recognize_google(audio)
        print("> {0}".format(prompt))

        #Remove pleasantries
        while "please" in prompt:
            prompt = prompt[:prompt.index("please")]
        while "Jeeves" in prompt:
            prompt = prompt[:prompt.index("Jeeves")]
        try:
            while prompt[-1] == " ":  #Remove possible extra whitespace on end to trim for command/content
                prompt = prompt[:-1]
        except:
            pass

        if "weather" in prompt:
            command = "weather"
        elif "show me files" in prompt:
            command = "filelist"
        elif "song" in prompt or "music" in prompt:
            command = "music"
        elif "summary" in prompt or "summarize" in prompt or "summarise" in prompt:
            command = "define"
        elif "shut" in prompt and "down" in prompt:
            command = "shutdown"
        elif "star trek" in prompt or "Star Trek" in prompt or "star Trek" in prompt or "Star trek" in prompt:
            command = "episode"
        elif "cast" in prompt:
            command = "cast"
        elif "headlines" in prompt or "headline" in prompt or "news" in prompt:
            command = "news"
        else:
            try:
                if prompt[:10] == "take notes":
                    command = "take notes"
                else:
                    command = prompt[:prompt.index(" ")]
            except:
                command = prompt
    
        if command == "open":
            app = prompt[prompt.index(command)+5:]

            print("Very well sir, I'll try and open {0} for you now...".format(app))
            tts.speak("Very well sir, I'll try and open {0} for you now...".format(app))

            try:
                if app == "Microsoft teams": #Specific fix for unrecognized name
                    app = "Microsoft Teams"
                    Commands.openApp(app,"Teams")
                else:
                    Commands.openApp(app)
            except FileNotFoundError:
                print("I'm afraid I couldn't find the file sir. Would you like to try again?")
                tts.speak("I'm afraid I couldn't find the file sir. Would you like to try again?")


        elif command == "take notes":
            r.pause_threshold = 4  #Max time in seconds to allow for pauses without stopping recording

            print("Alright then sir, go ahead, my pen is ready.")
            tts.speak("Alright then sir, go ahead, my pen is ready.")

            with mic as source:
                audio = r.listen(source)

            try:
                notes = r.recognize_google(audio)
                Commands.dictate(notes)
            except:
                print("Sorry sir, I didn't quite gather anything there. Try and run me again.")
                tts.speak("Sorry sir, I didn't quite gather anything there. Try and run me again.")


        elif command == "weather":
            if "1-hour" in prompt or "an hour" in prompt:
                try:
                    Commands.forecast(1)
                except:
                    print("Sorry sir, I couldn't fetch any data for that time. Perhaps try a different forecast.")
                    tts.speak("Sorry sir, I couldn't fetch any data for that time. Perhaps try a different forecast.")
    
            elif "2 hours" in prompt:
                try:
                    Commands.forecast(2)
                except:
                    print("Sorry sir, I couldn't fetch any data for that time. Perhaps try a different forecast.")
                    tts.speak("Sorry sir, I couldn't fetch any data for that time. Perhaps try a different forecast.")
    
            elif "3 hours" in prompt:
                try:
                    Commands.forecast(3)
                except:
                    print("Sorry sir, I couldn't fetch any data for that time. Perhaps try a different forecast.")
                    tts.speak("Sorry sir, I couldn't fetch any data for that time. Perhaps try a different forecast.")
    
            else:
                try:
                    Commands.weather()
                except:
                    print("Sorry sir, I couldn't fetch any data for that time. Perhaps try a different forecast.")
                    tts.speak("Sorry sir, I couldn't fetch any data for that time. Perhaps try a different forecast.")
    

        elif command == "filelist":
            if "under the name" in prompt:  #Correct way for this program to refer to the name of the files
                loc = prompt.index("under the name")
                name = prompt[loc+15:]

                try:
                    Commands.filelist("..",name)
                except:
                    print("I've run into a problem with the name you gave me sir. Please try again.")
                    tts.speak("I've run into a problem with the name you gave me sir. Please try again.")
    
            else:
                print("Sorry sir, I need you to phrase it as 'Show me files under the name...'.")
                tts.speak("Sorry sir, I need you to phrase it as 'Show me files under the name...'.")

    
        elif command == "repeat":
            while "this for me" in prompt:
                prompt = prompt.replace("this for me ", "")
            phrase = prompt[7:]
            print("I believe sir, you said '{0}'. Was I correct?".format(phrase))
            tts.speak("I believe sir, you said '{0}'. Was I correct?".format(phrase))

        elif command == "music":
            #Check Spotify API token is valid
            while str(requests.get("https://api.spotify.com/v1/playlists/566CNyS94IjywKSys66FJv", headers = {"Authorization": "Bearer "+Commands.access})) != "<Response [200]>":
                Commands.refreshSpotifyToken(spotifyRefreshToken,spotifyRefreshHeader)
                print("> Spotify Token Refreshed.")
                continue

            if "skip" in prompt or "next" in prompt:
                #Next Song Post Request
                requests.post("https://api.spotify.com/v1/me/player/next",headers = {"Authorization": "Bearer "+Commands.access})

            if "previous" in prompt or "last" in prompt:
                #Previous Song Post Request
                requests.post("https://api.spotify.com/v1/me/player/previous",headers = {"Authorization": "Bearer "+Commands.access})
            
            elif "pause" in prompt:
                #Pause Song Put Request
                requests.put("https://api.spotify.com/v1/me/player/pause", headers = {"Authorization": "Bearer "+Commands.access})
            
            elif "play" in prompt:
                #Play Song Put Request
                requests.put("https://api.spotify.com/v1/me/player/play", headers = {"Authorization": "Bearer "+Commands.access})
        
            elif "half volume" in prompt or "50%" in prompt:
                #Half volume Put Request
                requests.put("https://api.spotify.com/v1/me/player/volume", headers = {"Authorization": "Bearer "+Commands.access}, params = {"volume_percent":50})
            
            elif "full volume" in prompt or "100%" in prompt:
                #Full volume Put Request
                requests.put("https://api.spotify.com/v1/me/player/volume", headers = {"Authorization": "Bearer "+Commands.access}, params = {"volume_percent":100})
            
            elif "shuffle" in prompt:
                #Shuffle music Put Request
                requests.put("https://api.spotify.com/v1/me/player/shuffle",headers = {"Authorization": "Bearer "+Commands.access}, params = {"state":"true"})

        elif command == "define":
            if "summarize" in prompt:
                topic = prompt[prompt.index("summarize")+10:]
            elif "summarise" in prompt:
                topic = prompt[prompt.index("summarise")+10:]
            elif "summary" in prompt:
                topic = prompt[prompt.index("summary")+11:]

            if "sentences" in prompt:
                if "three" in prompt or "3" in prompt:
                    s = 3
                elif "two" in prompt or "2" in prompt:
                    s = 2
                elif "one" in prompt or "1" in prompt:
                    s = 1
            elif "sentence" in prompt:
                if "three" in prompt or "3" in prompt:
                    s = 3
                elif "two" in prompt or "2" in prompt:
                    s = 2
                elif "one" in prompt or "1" in prompt:
                    s = 1
            else:
                try:
                    s = int(prompt[prompt.index("sentence")-2])
                except:
                    pass

            topic[0].capitalize()

            try:
                Commands.define(topic,s)
            except:
                Commands.define(topic)

        elif command == "shutdown":
            print("Alright then sir, enjoy the rest of your day!")
            tts.speak("Alright then sir, enjoy the rest of your day!")
            time.sleep(1)
            playsound.playsound("jeeves_shutdown.mp3")
            attending = False
            break

        elif command == "episode":
            print("I shall check for you now sir.")
            tts.speak("I shall check for you now sir.")
            try:
                Commands.episode()
            except:
                print("Sorry sir, I seem to have run into some issues. Perhaps you'll have to try again.")
                tts.speak("Sorry sir, I seem to have run into some issues. Perhaps you'll have to try again.")

        elif command == "cast":
            name = prompt[prompt.find("cast")+8:]
            try:
                Commands.cast(name)
            except:
                print("Sorry sir, I was unable to fetch a cast for {0}.".format(name))
                tts.speak("Sorry sir, I was unable to fetch a cast for {0}.".format(name))

        elif command == "news":
            try:
                Commands.news()
            except:
                tts.speak("Sorry sir, I'm unable find any headlines at the moment. Perhaps try again later.")

        else:
            print("I'm sorry sir, did you say '{0}'?".format(command))
            print("Please go ahead and run me again, my hearing might be off.")
            tts.speak("I'm sorry sir, did you say '{0}'? Please go ahead and run me again, my hearing might be off.".format(command))
    
        caught = "" #Sound caught every two seconds on sleep cycle
        i = 0
        while i < 600:
            try:
                with mic as source: #Listen every two seconds
                    audio = r.listen(source, timeout=2.5, phrase_time_limit=2.5)
            except:
                continue
            try:
                caught = r.recognize_google(audio) #Try to analyse
            except:
                pass
        
            try: #If Jeeves is called in the two second cycle, trigger callback
                if "Jeeves" in caught or "geeves" in caught or "Jesus" in caught or "juice" in caught or "cheese" in caught:
                    callback = True
                    break
                elif ("skip" in caught or "next" in caught) and ("music" in caught or "song" in caught):
                    while str(requests.get("https://api.spotify.com/v1/playlists/566CNyS94IjywKSys66FJv", headers = {"Authorization": "Bearer "+Commands.access})) != "<Response [200]>":
                        Commands.refreshSpotifyToken(spotifyRefreshToken,spotifyRefreshHeader)
                        print("> Spotify Token Refreshed.")
                        continue
                    #Next Song Post Request
                    requests.post("https://api.spotify.com/v1/me/player/next",headers = {"Authorization": "Bearer "+Commands.access})
                    caught = ""
                    pass
            except:
                pass
            i += 1

        #If callback triggered, restart Jeeves loop
        if callback == True:
            continue
            
        print("I'm shutting down now sir.")
        tts.speak("I'm shutting down now sir.")
        attending = False

device = ["f21be663b348225da2fdb929b4f3ce814164db81"]